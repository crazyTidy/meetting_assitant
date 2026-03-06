"""DOCX generator for meeting minutes with Word heading styles."""
import logging
import re
from io import BytesIO
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


class DocxGenerator:
    """
    Generator for meeting minutes DOCX files following Chinese official document standards.

    Maps markdown headings to official document font specifications:
    - # (一级标题) → 二号小标宋 (22pt) - 公文大标题
    - ## (二级标题) → 三号黑体 (16pt) - 正文一级标题（一、）
    - ### (三级标题) → 三号楷体GB2312 (16pt) - 正文二级标题（（一））
    - #### (四级标题) → 三号仿宋_GB2312加粗 (16pt) - 正文三级标题（1.）
    - #####、###### (五级、六级标题) - 不使用（公文标准不允许）

    Page settings:
    - Paper size: A4 (210mm x 297mm)
    - Margins: 2.54cm (1 inch) all sides
    - Line spacing: 28pt fixed
    """

    # Font sizes
    FONT_SIZE_ERHAO = Pt(22)  # 二号 ~22pt
    FONT_SIZE_SANHAO = Pt(16)  # 三号 16pt

    # Line spacing: 28pt
    LINE_SPACING = Pt(28)

    # Margins: 2.54cm = 1 inch
    MARGIN_INCHES = 1.0

    # Paragraph spacing - reduced to minimize blank lines
    SPACING_BEFORE = Pt(0)
    SPACING_AFTER = Pt(0)

    def __init__(self):
        """Initialize the DOCX generator."""
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """Configure page settings for A4 paper with 2.54cm margins."""
        sections = self.doc.sections
        for section in sections:
            # Page dimensions for A4
            section.page_height = Inches(11.69)  # 297mm
            section.page_width = Inches(8.27)     # 210mm

            # Margins: 2.54cm = 1 inch
            section.top_margin = Inches(self.MARGIN_INCHES)
            section.bottom_margin = Inches(self.MARGIN_INCHES)
            section.left_margin = Inches(self.MARGIN_INCHES)
            section.right_margin = Inches(self.MARGIN_INCHES)

    def _set_font(self, run, font_name: str, size: Pt = FONT_SIZE_SANHAO, bold: bool = False):
        """
        Set font for a text run with Chinese font support.

        Args:
            run: The text run to apply font to
            font_name: Font name (e.g., '黑体', '楷体GB2312', '仿宋_GB2312')
            size: Font size
            bold: Whether text is bold
        """
        run.font.size = size
        run.font.bold = bold

        # Set ASCII font
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _add_paragraph(
        self,
        text: str = "",
        font_name: str = "仿宋_GB2312",
        alignment: WD_PARAGRAPH_ALIGNMENT = WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        first_line_indent: bool = True,
        bold: bool = False
    ) -> None:
        """
        Add a paragraph with standard formatting.

        Args:
            text: Paragraph text
            font_name: Font name
            alignment: Paragraph alignment
            first_line_indent: Whether to indent first line (2 characters)
            bold: Whether text is bold
        """
        para = self.doc.add_paragraph()
        para.alignment = alignment

        # Set paragraph spacing
        para.paragraph_format.line_spacing = self.LINE_SPACING
        para.paragraph_format.space_before = self.SPACING_BEFORE
        para.paragraph_format.space_after = self.SPACING_AFTER

        # Set first line indent (2 characters ≈ 0.8cm)
        if first_line_indent and alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
            para.paragraph_format.first_line_indent = Inches(0.32)

        if text:
            run = para.add_run(text)
            self._set_font(run, font_name, self.FONT_SIZE_SANHAO, bold)

    def _add_heading_by_level(self, text: str, level: int) -> None:
        """
        Add a heading according to Chinese official document standards.

        Args:
            text: Heading text
            level: Markdown heading level (1-4)
                   1=# 一级标题, 2=## 二级标题, 3=### 三级标题, 4=#### 四级标题
        """
        # Clamp to valid range (1-4) - 五级、六级标题不允许
        level = max(1, min(level, 4))

        para = self.doc.add_paragraph()
        para.paragraph_format.line_spacing = self.LINE_SPACING
        para.paragraph_format.first_line_indent = 0

        if level == 1:
            # 一级标题：二号小标宋，居中
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(12)
            run = para.add_run(text)
            self._set_font(run, "方正小标宋简体", self.FONT_SIZE_ERHAO, bold=True)

        elif level == 2:
            # 二级标题：三号黑体，左对齐
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            self._set_font(run, "黑体", self.FONT_SIZE_SANHAO, bold=True)

        elif level == 3:
            # 三级标题：三号楷体GB2312，左对齐
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            self._set_font(run, "楷体GB2312", self.FONT_SIZE_SANHAO, bold=False)

        else:  # level == 4
            # 四级标题：三号仿宋_GB2312加粗，左对齐
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO, bold=True)

    def add_title(self, meeting_title: str) -> None:
        """
        Add the main document title using 一级标题 style.

        Args:
            meeting_title: Title of the meeting
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.paragraph_format.line_spacing = self.LINE_SPACING
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.first_line_indent = 0

        run = para.add_run(f"{meeting_title}会议纪要")
        self._set_font(run, "方正小标宋简体", self.FONT_SIZE_ERHAO, bold=True)

    def parse_markdown_content(self, markdown: str) -> None:
        """
        Parse markdown content and add to document.

        Args:
            markdown: Markdown formatted content
        """
        lines = markdown.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # Empty line - skip it (no blank lines in official documents)
            if not line:
                i += 1
                continue

            # ATX style headings - map to official document standards
            # Only support # ## ### #### (level 1-4)
            # ##### and ###### are not allowed in official documents
            heading_match = re.match(r'^(#{1,4})\s+(.+)', line)
            if heading_match:
                marks = heading_match.group(1)
                heading_text = heading_match.group(2)
                level = len(marks)  # 1-4
                self._add_heading_by_level(heading_text, level)
                i += 1
                continue

            # Chinese numbered heading: （一）（二）（三） - 三级标题
            elif re.match(r'^[（\(][一二三四五六七八九十][）\)]', line):
                self._add_heading_by_level(line, level=3)
                i += 1
                continue

            # Chinese numbered heading variant: 一、二、三、 - 二级标题
            elif re.match(r'^[一二三四五六七八九十]+、', line):
                self._add_heading_by_level(line, level=2)
                i += 1
                continue

            # Bold heading: **标题** - 二级标题
            elif re.match(r'^\*{2}[^*]+\*{2}\s*$', line):
                heading_text = re.sub(r'^\*{2}|\*{2}\s*$', '', line)
                if heading_text:
                    self._add_heading_by_level(heading_text, level=2)
                i += 1
                continue

            # Numbered list: 1. 2. 3.
            elif re.match(r'^(\d+)\.\s+', line):
                match = re.match(r'^(\d+)\.\s+(.+)', line)
                if match:
                    num, content = match.groups()
                    para = self.doc.add_paragraph()
                    para.paragraph_format.line_spacing = self.LINE_SPACING
                    para.paragraph_format.space_before = self.SPACING_BEFORE
                    para.paragraph_format.space_after = self.SPACING_AFTER
                    para.paragraph_format.first_line_indent = 0

                    # Number part - use bold for 四级标题 style
                    run = para.add_run(f"{num}. ")
                    self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO, bold=True)

                    # Content part
                    self._parse_inline_text(para, content)
                i += 1
                continue

            # Lettered list: a. b. c. or A. B. C.
            elif re.match(r'^([a-zA-Z])\.\s+', line):
                match = re.match(r'^([a-zA-Z])\.\s+(.+)', line)
                if match:
                    letter, content = match.groups()
                    para = self.doc.add_paragraph()
                    para.paragraph_format.line_spacing = self.LINE_SPACING
                    para.paragraph_format.space_before = self.SPACING_BEFORE
                    para.paragraph_format.space_after = self.SPACING_AFTER
                    para.paragraph_format.first_line_indent = 0

                    # Letter part
                    run = para.add_run(f"{letter}. ")
                    self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO, bold=False)

                    # Content part
                    self._parse_inline_text(para, content)
                i += 1
                continue

            # Bullet points: - or * (but not - - - which is horizontal rule)
            elif re.match(r'^[\-*]\s+', line):
                text = line[2:] if len(line) > 2 else ''
                para = self.doc.add_paragraph()
                para.paragraph_format.line_spacing = self.LINE_SPACING
                para.paragraph_format.space_before = self.SPACING_BEFORE
                para.paragraph_format.space_after = self.SPACING_AFTER
                para.paragraph_format.first_line_indent = Inches(0.32)

                # Bullet symbol
                run = para.add_run("• ")
                self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO)

                # Content
                if text:
                    self._parse_inline_text(para, text)
                i += 1
                continue

            # Table rows: | col1 | col2 | col3 | (must check before horizontal rule)
            elif line.startswith('|') and line.endswith('|'):
                cells = [c.strip() for c in line.split('|')[1:-1] if c.strip()]
                if cells:
                    # Check if this is a separator row (contains only -, =, or :)
                    is_separator = all(re.match(r'^[:\\-]+$', cell) for cell in cells)
                    if is_separator:
                        # Skip separator row
                        i += 1
                        continue

                    para = self.doc.add_paragraph()
                    para.paragraph_format.line_spacing = self.LINE_SPACING
                    text = '   |   '.join(cells)
                    run = para.add_run(text)
                    self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO, bold=False)
                i += 1
                continue

            # Horizontal rule: --- or *** (but not inside table)
            elif line.strip() in ('---', '***', '___'):
                # Add a horizontal line using border
                para = self.doc.add_paragraph()
                para.paragraph_format.line_spacing = self.LINE_SPACING
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                # Add a simple line using paragraph border
                run = para.add_run("─" * 50)
                self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO)
                i += 1
                continue

            # Code block: ``` or ~~~
            elif line.startswith('```') or line.startswith('~~~'):
                # Find end of code block
                lang = line[3:].strip() if line.startswith('```') else ''
                code_lines = []
                i += 1
                while i < len(lines):
                    if lines[i].startswith('```') or lines[i].startswith('~~~'):
                        break
                    code_lines.append(lines[i])
                    i += 1

                # Add code as preformatted text
                para = self.doc.add_paragraph()
                para.paragraph_format.line_spacing = Pt(12)  # Tighter spacing for code
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.first_line_indent = 0

                for code_line in code_lines:
                    if code_line:
                        run = para.add_run(code_line + '\n')
                        self._set_font(run, "Courier New", Pt(11), bold=False)
                    else:
                        para.add_run('\n')

                i += 1  # Skip closing ```
                continue

            # Blockquote: > text
            elif line.startswith('> '):
                text = line[2:]
                para = self.doc.add_paragraph()
                para.paragraph_format.line_spacing = self.LINE_SPACING
                para.paragraph_format.space_before = self.SPACING_BEFORE
                para.paragraph_format.space_after = self.SPACING_AFTER
                para.paragraph_format.first_line_indent = Inches(0.32)
                para.paragraph_format.left_indent = Inches(0.5)

                # Add quote marker
                run = para.add_run("▌ ")
                self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO)

                # Content
                self._parse_inline_text(para, text)
                i += 1
                continue

            # Regular paragraph
            else:
                para = self.doc.add_paragraph()
                para.paragraph_format.line_spacing = self.LINE_SPACING
                para.paragraph_format.space_before = self.SPACING_BEFORE
                para.paragraph_format.space_after = self.SPACING_AFTER
                para.paragraph_format.first_line_indent = Inches(0.32)
                self._parse_inline_text(para, line)
                i += 1
                continue

    def _parse_inline_text(self, paragraph, text: str) -> None:
        """
        Parse inline markdown formatting and add to paragraph.
        Strips markdown markers and outputs plain text.

        Args:
            paragraph: The paragraph to add text to
            text: Text with potential inline formatting
        """
        if not text:
            return

        # Remove all markdown formatting and add as plain text
        # Process code spans first
        clean_text = re.sub(r'`([^`]+)`', r'\1', text)

        # Remove bold+italic: ***text*** or ___text___
        clean_text = re.sub(r'\*\*\*([^*]+)\*\*\*', r'\1', clean_text)
        clean_text = re.sub(r'___([^_]+)___', r'\1', clean_text)

        # Remove bold: **text** or __text__
        clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean_text)
        clean_text = re.sub(r'__([^_]+)__', r'\1', clean_text)

        # Remove italic: *text* or _text_
        clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)
        clean_text = re.sub(r'_([^_]+)_', r'\1', clean_text)

        # Remove strikethrough: ~~text~~
        clean_text = re.sub(r'~~([^~]+)~~', r'\1', clean_text)

        # Remove links: [text](url) -> text
        clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_text)

        # Add the cleaned text as a single run
        if clean_text:
            run = paragraph.add_run(clean_text)
            self._set_font(run, "仿宋_GB2312", self.FONT_SIZE_SANHAO, bold=False)

    def generate(self, meeting_title: str, content: str) -> bytes:
        """
        Generate the complete DOCX document.

        Args:
            meeting_title: Title of the meeting
            content: Markdown formatted content

        Returns:
            DOCX file as bytes
        """
        logger.info(f"[DOCX] Generating document for meeting: {meeting_title}")

        # Add title
        self.add_title(meeting_title)

        # Parse and add content
        self.parse_markdown_content(content)

        # Save to bytes
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)

        logger.info(f"[DOCX] Document generated successfully, size: {len(buffer.getvalue())} bytes")

        return buffer.getvalue()


def generate_meeting_minutes_docx(meeting_title: str, content: str) -> bytes:
    """
    Convenience function to generate meeting minutes DOCX.

    Args:
        meeting_title: Title of the meeting
        content: Markdown formatted content

    Returns:
        DOCX file as bytes
    """
    generator = DocxGenerator()
    return generator.generate(meeting_title, content)
